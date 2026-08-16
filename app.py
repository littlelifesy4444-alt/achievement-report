# -*- coding: utf-8 -*-
from __future__ import annotations

import base64
import io
import json
import mimetypes
import os
import re
import zipfile
from collections import defaultdict
from typing import Literal

import streamlit as st
from openai import OpenAI
from pydantic import BaseModel, Field
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak

pdfmetrics.registerFont(UnicodeCIDFont("HYSMyeongJo-Medium"))

MODEL = "gpt-5.6"

class MemoryUploadedFile:
    """ZIP 안의 파일을 Streamlit UploadedFile처럼 다루기 위한 간단한 래퍼."""
    def __init__(self, name: str, data: bytes):
        self.name = name
        self._data = data

    def getvalue(self):
        return self._data


MANUAL_RULES = """
성취도평가 학생 전달용 리포트 제작 기준

1. 모든 문항을 문항번호별로 분석한다.
2. 각 문항에 '문항 분석 영역', '핵심 출제 포인트'를 작성한다.
   문항 분석 영역은 아래 11개 대분류 중 하나만 사용한다.
   - 문장의 형식·동사: 1~5형식, 목적격보어, 사역/지각동사, 동사별 문형, 가주어·문장구조
   - 시제: 현재/과거/미래, 진행, 완료, 시간 부사절
   - 조동사: can, may, must, should, have to, 조동사+완료형
   - 수동태: be p.p. 및 수동 관련 표현
   - 준동사: to부정사, 동명사, 현재분사/과거분사, 완료부정사
   - 품사·어휘: 품사 구별, 접속사, 어휘 결합 및 일반 어법
   - 관계사: 관계대명사, 관계부사 및 관계사 구문
   - 전치사: 전치사 선택과 전치사 결합 표현
   - 비교 표현: 원급, 비교급, 최상급 및 비교 구문
   - 도치·특수구문·화법: 도치, 강조, 생략, 병렬, 일치, 직접/간접화법 등
   - 기타 문법·구문: 위 영역으로 자연스럽게 분류되지 않는 문법
3. 공식 정답표가 있으면 반드시 우선 적용한다.
4. 공식 정답표가 없으면 시험지를 분석해 정답을 판단하되, 정답이 모호하거나 복수 정답 가능성이 있으면 추측하지 않고
   '확인 필요' 또는 '채점 제외'로 처리한다.
5. 객관식은 정답과 학생 답안을 비교해 ○/×로 판정한다.
6. 서술형은 확실한 채점 기준이 있는 경우만 채점한다. 확실하지 않으면 임의 채점하지 않는다.
7. 채점 제외/확인 필요 문항은 정답률 계산의 분모에서도 제외한다.
8. 문항별 리포트는 6열 구조를 사용한다:
   번호 / 문항 분석 / 핵심 출제 포인트 / 정오 / 분석 근거 / 학생 진단
9. 번호 열과 정오 열은 매우 좁게 사용하고, 설명 열을 넓게 사용한다.
10. 맞은 문항의 학생 진단은 공란으로 둔다.
11. 틀린 문항의 학생 진단은 정답 자체를 알려주기보다 '시제 구별 보완 필요', '목적격보어 형태 확인 필요'처럼
    짧고 구체적인 학습 포인트를 제시한다.
12. 객관식 문항도 '분석 근거'를 표시한다. 단, 정답 근거를 짧고 명확하게 작성하고 문제 전체 문장을 재현하지 않는다.
13. 서술형 분석 근거는 전체 정답 문장이 아니라 채점에 필요한 부분답만 짧게 표시한다.
14. 서술형의 전체 기준답안과 학생이 작성한 전체 답안은 리포트에 공개하지 않는다.
16. 틀린 서술형에만 짧은 오류 진단을 작성한다.
17. 영역별 표에는 영역 / 정답 / 채점 문항 / 정답률 / 판정을 표시한다.
18. 강점과 보완 영역은 실제 정답률과 해당 영역의 채점 문항 수를 함께 고려한다.
   채점 문항이 1~2개뿐인 영역은 '참고'로 표시하고 강점/보완 영역 자동 선정에서는 우선 제외한다.
19. 오답이 집중된 영역과 상대적으로 안정적인 영역을 실제 계산 결과에서 찾는다.
20. 종합 코멘트는 '현재 성취 수준 → 강점 → 보완 영역 → 구체적 복습 방향' 순으로 작성한다.
21. 학생의 태도, 노력, 성격, 수업 참여도 등 답안지에서 확인할 수 없는 내용은 추정하지 않는다.
22. 학생 간 비교나 서열 표현은 사용하지 않는다.
23. Word 디자인:
    - 표 헤더: 진한 파란색 + 흰 글씨
    - 본문: 옅은 교차 행 음영
    - 정오: ○ 연한 초록 / × 연한 주황·붉은색 / 채점 제외·확인 필요 연한 노랑
    - 오답 학생 진단: 연한 노랑
    - 서술형 분석 근거: 옅은 회색·파랑 계열
    - 번호와 정오 열은 매우 좁고 설명 열은 넓게
24. 최종 구성:
    1) 문항별 분석 및 정오표
    2) 영역별 성취도
    3) 오답 기반 진단
    4) 종합 학습 코멘트
"""

class QuestionAnalysis(BaseModel):
    number: int
    question_type: Literal["objective", "subjective"]
    points: float = Field(description="문항 배점. 시험지에서 확인 불가하면 0.")
    area: Literal["문장의 형식·동사", "시제", "조동사", "수동태", "준동사", "품사·어휘", "관계사", "전치사", "비교 표현", "도치·특수구문·화법", "기타 문법·구문"]
    point: str
    answer_status: Literal["certain", "confirm", "exclude"]
    correct_answer: str = Field(description="객관식 정답 번호 또는 서술형 채점에 필요한 부분답만. 전체 정답문장 금지.")
    evidence: str = Field(description="서술형 리포트에 쓸 짧은 부분답 근거. 객관식은 내부 분석용으로만 사용 가능.")
    wrong_diagnosis: str = Field(description="이 문항을 틀렸을 때 쓸 짧은 학습 진단.")

class ExamAnalysis(BaseModel):
    title: str
    stated_total_points: float = Field(description="시험지에 적힌 총점. 확인 불가하면 0.")
    total_questions: int
    notes: str
    questions: list[QuestionAnalysis]

class GradeItem(BaseModel):
    number: int
    result: Literal["correct", "wrong", "confirm", "exclude"]
    evidence: str = Field(description="서술형 채점에 필요한 짧은 부분 근거. 학생 전체 답안 재현 금지.")
    diagnosis: str = Field(description="wrong일 때만 짧고 구체적 진단. correct일 때 빈 문자열.")

class StudentGrading(BaseModel):
    student_name: str
    class_name: str
    items: list[GradeItem]


def client_for(api_key: str) -> OpenAI:
    return OpenAI(api_key=api_key)


def file_item(uploaded, *, detail="high"):
    data = uploaded.getvalue()
    name = uploaded.name
    suffix = os.path.splitext(name)[1].lower()
    b64 = base64.b64encode(data).decode("ascii")

    if suffix == ".pdf":
        return {
            "type": "input_file",
            "filename": name,
            "file_data": f"data:application/pdf;base64,{b64}",
            "detail": detail,
        }

    if suffix in (".png", ".jpg", ".jpeg", ".webp"):
        mime = mimetypes.guess_type(name)[0] or "image/png"
        return {
            "type": "input_image",
            "image_url": f"data:{mime};base64,{b64}",
            "detail": "high",
        }

    raise ValueError(f"지원하지 않는 파일 형식: {name}")


def analyze_exam(api_key: str, exam_file, answer_key_file=None) -> ExamAnalysis:
    client = client_for(api_key)
    content = [
        {
            "type": "input_text",
            "text": (
                "당신은 중학교 영어 성취도평가 리포트의 시험 분석 담당자입니다.\n"
                "아래 제작 기준을 반드시 따르세요.\n\n"
                + MANUAL_RULES
                + "\n\n시험지의 모든 문항을 빠짐없이 번호별로 분석하세요. "
                  "문제 자체의 오류, 복수 정답 가능성, 정답 불확실성이 있으면 억지로 정답을 확정하지 마세요. "
                  "문항 배점도 시험지에서 확인하세요. "
                  "서술형 correct_answer/evidence에는 전체 정답 문장을 넣지 말고 채점에 필요한 부분만 넣으세요."
            ),
        },
        file_item(exam_file, detail="high"),
    ]
    if answer_key_file is not None:
        content += [
            {
                "type": "input_text",
                "text": "다음 파일은 공식 정답표 또는 공식 채점기준입니다. 시험지 자체 분석보다 반드시 우선 적용하세요.",
            },
            file_item(answer_key_file, detail="high"),
        ]

    response = client.responses.parse(
        model=MODEL,
        reasoning={"effort": "medium"},
        input=[{"role": "user", "content": content}],
        text_format=ExamAnalysis,
        store=False,
    )
    if response.output_parsed is None:
        raise RuntimeError("시험지 분석 결과를 구조화하지 못했습니다.")
    return response.output_parsed


def grade_student(api_key: str, exam: ExamAnalysis, answer_file) -> StudentGrading:
    client = client_for(api_key)
    exam_json = exam.model_dump_json(indent=2)

    prompt = (
        "당신은 학생 답안 채점 담당자입니다. 아래 시험 분석값과 첨부 학생 답안을 비교해 모든 문항을 채점하세요.\n\n"
        + MANUAL_RULES
        + "\n\n시험 분석 JSON:\n"
        + exam_json
        + "\n\n규칙:\n"
          "- 시험 분석에서 answer_status=exclude인 문항은 무조건 exclude.\n"
          "- answer_status=confirm인 문항은 충분한 근거가 없으면 confirm 유지.\n"
          "- 학생 답안이 안 보이거나 판독이 불확실하면 추측하지 말고 confirm.\n"
          "- correct 문항의 diagnosis는 반드시 빈 문자열.\n"
          "- wrong 문항의 diagnosis는 짧고 구체적인 학습 포인트.\n"
          "- 객관식 evidence도 정답 판단 근거를 짧게 작성한다. 문제 전체 문장을 복사하지 않는다.\n"
          "- 서술형 evidence에 학생의 전체 답안을 옮기지 말고 채점에 필요한 부분만 표시.\n"
          "- 학생 이름/반은 답안지와 파일명에서 확인되는 정보만 사용."
    )

    response = client.responses.parse(
        model=MODEL,
        reasoning={"effort": "medium"},
        input=[
            {
                "role": "user",
                "content": [
                    {"type": "input_text", "text": prompt},
                    file_item(answer_file, detail="high"),
                ],
            }
        ],
        text_format=StudentGrading,
        store=False,
    )
    if response.output_parsed is None:
        raise RuntimeError(f"{answer_file.name}: 채점 결과를 구조화하지 못했습니다.")

    result = response.output_parsed
    if not result.student_name.strip():
        result.student_name = os.path.splitext(answer_file.name)[0]
    return result


def judgment(rate: float) -> str:
    if rate >= 90:
        return "매우 우수"
    if rate >= 80:
        return "우수"
    if rate >= 70:
        return "보통"
    if rate >= 60:
        return "보완"
    return "우선 보완"


def calculate(exam: ExamAnalysis, grading: StudentGrading):
    qmap = {q.number: q for q in exam.questions}
    imap = {x.number: x for x in grading.items}

    rows = []
    areas = defaultdict(lambda: {"correct": 0, "count": 0, "score": 0.0, "possible": 0.0})
    correct = wrong = graded = excluded = confirm = 0
    score = possible = 0.0

    for n in sorted(qmap):
        q = qmap[n]
        item = imap.get(n)

        if q.answer_status == "exclude":
            mark = "제외"
        elif item is None:
            mark = "확인"
        else:
            mark = {
                "correct": "○",
                "wrong": "×",
                "confirm": "확인",
                "exclude": "제외",
            }.get(item.result, "확인")

        if mark == "○":
            correct += 1
            graded += 1
        elif mark == "×":
            wrong += 1
            graded += 1
        elif mark == "제외":
            excluded += 1
        else:
            confirm += 1

        if mark in ("○", "×"):
            pts = max(float(q.points or 0), 0.0)
            possible += pts
            areas[q.area]["count"] += 1
            areas[q.area]["possible"] += pts
            if mark == "○":
                score += pts
                areas[q.area]["correct"] += 1
                areas[q.area]["score"] += pts

        # 객관식/서술형 모두 분석근거 표시.
        # 서술형은 전체 답안이 아닌 채점에 필요한 부분답만 사용한다.
        evidence = q.evidence
        if item and item.evidence.strip():
            evidence = item.evidence.strip()

        diagnosis = q.wrong_diagnosis
        if item and item.diagnosis.strip():
            diagnosis = item.diagnosis.strip()
        if mark == "○":
            diagnosis = "정답"

        rows.append(
            {
                "number": n,
                "area": q.area,
                "point": q.point,
                "mark": mark,
                "evidence": evidence,
                "diagnosis": diagnosis,
                "question_type": q.question_type,
            }
        )

    rate = round(correct / graded * 100, 1) if graded else 0.0
    area_rows = []
    for area, s in areas.items():
        ar = round(s["correct"] / s["count"] * 100, 1) if s["count"] else 0.0
        area_rows.append(
            {
                "area": area,
                "correct": s["correct"],
                "count": s["count"],
                "rate": ar,
                "judgment": "참고" if s["count"] < 3 else judgment(ar),
            }
        )
    area_rows.sort(key=lambda x: (x["rate"], -x["count"], x["area"]))

    return {
        "name": grading.student_name.strip() or "학생",
        "class": grading.class_name.strip(),
        "correct": correct,
        "wrong": wrong,
        "graded": graded,
        "excluded": excluded,
        "confirm": confirm,
        "rate": rate,
        "score": round(score, 1),
        "possible": round(possible, 1),
        "rows": rows,
        "areas": area_rows,
    }


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_text(cell, text, *, bold=False, color=None, align=None, size=7.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    if align is not None:
        p.alignment = align
    r = p.add_run(str(text))
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(size)
    r.bold = bold
    if color is not None:
        r.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement("w:tblHeader")
    h.set(qn("w:val"), "true")
    trPr.append(h)


def set_column_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(31, 78, 121)
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def comment_text(result):
    if not result["areas"]:
        return "채점 가능한 문항이 없어 종합 학습 코멘트를 산출할 수 없습니다."

    reliable = [a for a in result["areas"] if a["count"] >= 3]
    pool = reliable if reliable else result["areas"]

    best = max(pool, key=lambda x: (x["rate"], x["count"]))
    weak = min(pool, key=lambda x: (x["rate"], -x["count"]))

    diag_counts = defaultdict(int)
    for row in result["rows"]:
        if row["mark"] == "×" and row["diagnosis"]:
            diag_counts[row["diagnosis"]] += 1

    review = ""
    if diag_counts:
        ranked = sorted(diag_counts.items(), key=lambda x: (-x[1], x[0]))[:3]
        review = " 우선 복습할 내용은 " + ", ".join(
            f"{d}({c}문항)" for d, c in ranked
        ) + "입니다."

    note = ""
    if not reliable:
        note = " 영역별 문항 수가 적어 강점·보완 영역 해석은 참고용입니다."

    return (
        f'{result["name"]} 학생은 채점 가능한 {result["graded"]}문항 중 '
        f'{result["correct"]}문항을 맞혀 정답률 {result["rate"]}%를 기록했습니다. '
        f'상대적으로 안정적인 영역은 {best["area"]}({best["correct"]}/{best["count"]}, {best["rate"]}%)이며, '
        f'우선 보완이 필요한 영역은 {weak["area"]}({weak["correct"]}/{weak["count"]}, {weak["rate"]}%)입니다.'
        + review + note
    )

def make_docx(exam: ExamAnalysis, result) -> bytes:
    doc = Document()
    sec = doc.sections[0]
    # A4 가로 방향: 긴 설명 열의 가독성을 확보
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.top_margin = Inches(0.38)
    sec.bottom_margin = Inches(0.38)
    sec.left_margin = Inches(0.35)
    sec.right_margin = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(9.2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(exam.title + "\n학생 학습 리포트")
    r.bold = True
    r.font.size = Pt(16.5)
    r.font.color.rgb = RGBColor(31, 78, 121)
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")

    summary = doc.add_table(rows=2, cols=7)
    summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["학생", "반", "채점 문항", "정답", "정답률", "채점 점수", "보완 영역"]
    reliable_areas = [a for a in result["areas"] if a["count"] >= 3]
    summary_pool = reliable_areas if reliable_areas else result["areas"]
    weak_area = min(summary_pool, key=lambda x: (x["rate"], -x["count"]))["area"] if summary_pool else "-"
    values = [
        result["name"],
        result["class"],
        result["graded"],
        result["correct"],
        f'{result["rate"]}%',
        f'{result["score"]}/{result["possible"]}',
        weak_area,
    ]
    for i, h in enumerate(headers):
        cell_text(summary.cell(0, i), h, bold=True, color=(255, 255, 255),
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=8.7)
        shade(summary.cell(0, i), "1F4E79")
        cell_text(summary.cell(1, i), values[i], bold=(i == 4),
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=9.0)
        if i == 4:
            shade(summary.cell(1, i), "D9EAD3")

    notes = []
    if result["excluded"]:
        notes.append(f'채점 제외 {result["excluded"]}문항')
    if result["confirm"]:
        notes.append(f'확인 필요 {result["confirm"]}문항')
    if notes:
        p = doc.add_paragraph()
        rr = p.add_run("※ " + ", ".join(notes) + "은 정답률과 채점 가능 점수 계산에서 제외했습니다.")
        rr.font.size = Pt(8.0)

    heading(doc, "1. 문항별 분석 및 정오표")
    tb = doc.add_table(rows=1, cols=6)
    tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    tb.autofit = False

    # 번호/정오 열은 매우 좁게, 설명 열은 넓게
    col_widths = [
        Inches(0.432),  # 번호 4%
        Inches(1.081),  # 문항 분석 10%
        Inches(3.135),  # 핵심 출제 포인트 29%
        Inches(0.541),  # 정오 5%
        Inches(3.243),  # 분석 근거 30%
        Inches(2.378),  # 학생 진단 22%
    ]
    set_column_widths(tb, col_widths)

    hdr = ["번호", "문항 분석", "핵심 출제 포인트", "정오", "분석 근거", "학생 진단"]
    for i, h in enumerate(hdr):
        cell_text(tb.cell(0, i), h, bold=True, color=(255, 255, 255),
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=8.8)
        shade(tb.cell(0, i), "1F4E79")
    repeat_header(tb.rows[0])

    for idx, rowdata in enumerate(result["rows"], 1):
        row = tb.add_row().cells
        # 객관식과 서술형 모두 분석근거를 표시
        evidence_to_show = rowdata["evidence"]
        vals = [
            rowdata["number"],
            rowdata["area"],
            rowdata["point"],
            rowdata["mark"],
            evidence_to_show,
            rowdata["diagnosis"],
        ]
        for j, v in enumerate(vals):
            cell_text(
                row[j],
                v,
                align=WD_ALIGN_PARAGRAPH.CENTER if j in (0, 3) else WD_ALIGN_PARAGRAPH.LEFT,
                size=8.8 if j not in (0,3) else 9.2,
            )
            row[j].width = col_widths[j]

        if idx % 2 == 0:
            for c in row:
                shade(c, "EEF3F8")
        if rowdata["mark"] == "○":
            shade(row[3], "D9EAD3")
        elif rowdata["mark"] == "×":
            shade(row[3], "FCE4D6")
            shade(row[5], "FFF2CC")
        else:
            shade(row[3], "FFF2CC")

        # 분석근거는 객관식/서술형 모두 표시하되, 행 전체 교차 음영을 유지한다.

    heading(doc, "2. 영역별 성취도")
    at = doc.add_table(rows=1, cols=5)
    at.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["영역", "정답", "채점 문항", "정답률", "판정"]):
        cell_text(at.cell(0, i), h, bold=True, color=(255, 255, 255),
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=8.5)
        shade(at.cell(0, i), "1F4E79")
    repeat_header(at.rows[0])

    for idx, a in enumerate(result["areas"], 1):
        row = at.add_row().cells
        vals = [a["area"], a["correct"], a["count"], f'{a["rate"]}%', a["judgment"]]
        for j, v in enumerate(vals):
            cell_text(row[j], v, align=WD_ALIGN_PARAGRAPH.CENTER if j else WD_ALIGN_PARAGRAPH.LEFT, size=8.5)
        if idx % 2 == 0:
            for c in row:
                shade(c, "EEF3F8")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    rr = p.add_run("※ 채점 문항이 1~2개인 영역은 표본이 적어 판정을 '참고'로 표시합니다.")
    rr.font.size = Pt(8.0)

    heading(doc, "3. 오답 기반 진단")
    diag_counts = defaultdict(int)
    for x in result["rows"]:
        if x["mark"] == "×" and x["diagnosis"]:
            diag_counts[x["diagnosis"]] += 1
    if diag_counts:
        ranked_diags = sorted(diag_counts.items(), key=lambda x: (-x[1], x[0]))[:8]
        for d, count in ranked_diags:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.add_run(f"• {d} — 관련 오답 {count}문항")
    else:
        doc.add_paragraph("채점된 문항에서 별도 오답 진단 항목이 없습니다.")

    heading(doc, "4. 종합 학습 코멘트")
    doc.add_paragraph(comment_text(result))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()



def _pdf_p(text, style):
    """ReportLab Paragraph helper with safe text conversion."""
    s = str(text if text is not None else "")
    s = s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    s = s.replace("\n", "<br/>")
    return Paragraph(s, style)


def make_pdf(exam: ExamAnalysis, result) -> bytes:
    """Create a print-stable A4 landscape PDF directly, independent of Google Docs."""
    buf = io.BytesIO()
    page_w, page_h = landscape(A4)
    doc = SimpleDocTemplate(
        buf,
        pagesize=landscape(A4),
        leftMargin=0.30 * inch,
        rightMargin=0.30 * inch,
        topMargin=0.20 * inch,
        bottomMargin=0.20 * inch,
        title=f'{result["name"]} 학생 학습 리포트',
    )

    font = "HYSMyeongJo-Medium"
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "KTitle", parent=styles["Title"], fontName=font, fontSize=15,
        leading=16, alignment=TA_CENTER, textColor=colors.HexColor("#1F4E79"),
        spaceAfter=7,
    )
    heading_style = ParagraphStyle(
        "KHeading", parent=styles["Heading2"], fontName=font, fontSize=11,
        leading=11.5, textColor=colors.HexColor("#1F4E79"), spaceBefore=6, spaceAfter=4,
    )
    body = ParagraphStyle(
        "KBody", parent=styles["BodyText"], fontName=font, fontSize=7.9,
        leading=8.6, alignment=TA_LEFT,
    )
    body_center = ParagraphStyle(
        "KCenter", parent=body, alignment=TA_CENTER,
    )
    small = ParagraphStyle(
        "KSmall", parent=body, fontSize=7.2, leading=8.0,
    )
    header_style = ParagraphStyle(
        "KHeader", parent=body_center, fontSize=7.8, leading=8.4,
        textColor=colors.white,
    )

    story = []
    story.append(_pdf_p(exam.title + "<br/>학생 학습 리포트", title_style))

    reliable_areas = [a for a in result["areas"] if a["count"] >= 3]
    summary_pool = reliable_areas if reliable_areas else result["areas"]
    weak_area = min(summary_pool, key=lambda x: (x["rate"], -x["count"]))["area"] if summary_pool else "-"

    sh = ["학생", "반", "채점 문항", "정답", "정답률", "채점 점수", "보완 영역"]
    sv = [
        result["name"], result["class"], result["graded"], result["correct"],
        f'{result["rate"]}%', f'{result["score"]}/{result["possible"]}', weak_area
    ]
    summary_data = [
        [_pdf_p(x, header_style) for x in sh],
        [_pdf_p(x, body_center) for x in sv],
    ]
    usable = page_w - 0.60 * inch
    summary = Table(summary_data, colWidths=[usable / 7.0] * 7, repeatRows=1)
    summary.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#1F4E79")),
        ("BACKGROUND", (4,1), (4,1), colors.HexColor("#D9EAD3")),
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ("ALIGN", (0,0), (-1,-1), "CENTER"),
        ("GRID", (0,0), (-1,-1), 0.25, colors.HexColor("#B7C9D6")),
        ("LEFTPADDING", (0,0), (-1,-1), 2.5),
        ("RIGHTPADDING", (0,0), (-1,-1), 2.5),
        ("TOPPADDING", (0,0), (-1,-1), 1.2),
        ("BOTTOMPADDING", (0,0), (-1,-1), 1.2),
    ]))
    story.append(summary)

    notes = []
    if result["excluded"]:
        notes.append(f'채점 제외 {result["excluded"]}문항')
    if result["confirm"]:
        notes.append(f'확인 필요 {result["confirm"]}문항')
    if notes:
        story.append(Spacer(1, 1.5))
        story.append(_pdf_p("※ " + ", ".join(notes) + "은 정답률과 채점 가능 점수 계산에서 제외했습니다.", small))

    story.append(_pdf_p("1. 문항별 분석 및 정오표", heading_style))

    # Requested fixed proportions: 4 / 10 / 29 / 5 / 30 / 22 = 100%
    proportions = [0.04, 0.10, 0.29, 0.05, 0.30, 0.22]
    widths = [usable * p for p in proportions]
    headers = ["번호", "문항 분석", "핵심 출제 포인트", "정오", "분석 근거", "학생 진단"]
    data = [[_pdf_p(x, header_style) for x in headers]]

    for rowdata in result["rows"]:
        vals = [
            rowdata["number"], rowdata["area"], rowdata["point"], rowdata["mark"],
            rowdata["evidence"], rowdata["diagnosis"],
        ]
        data.append([
            _pdf_p(vals[0], body_center),
            _pdf_p(vals[1], body),
            _pdf_p(vals[2], body),
            _pdf_p(vals[3], body_center),
            _pdf_p(vals[4], body),
            _pdf_p(vals[5], body),
        ])

    table = Table(data, colWidths=widths, repeatRows=1, hAlign="CENTER")
    ts = [
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#1F4E79")),
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ("GRID", (0,0), (-1,-1), 0.25, colors.HexColor("#B7C9D6")),
        ("LEFTPADDING", (0,0), (-1,-1), 2.2),
        ("RIGHTPADDING", (0,0), (-1,-1), 2.2),
        ("TOPPADDING", (0,0), (-1,-1), 1.0),
        ("BOTTOMPADDING", (0,0), (-1,-1), 1.0),
    ]
    for i, rowdata in enumerate(result["rows"], start=1):
        if i % 2 == 0:
            ts.append(("BACKGROUND", (0,i), (-1,i), colors.HexColor("#EEF3F8")))
        if rowdata["mark"] == "○":
            ts.append(("BACKGROUND", (3,i), (3,i), colors.HexColor("#D9EAD3")))
        elif rowdata["mark"] == "×":
            ts.append(("BACKGROUND", (3,i), (3,i), colors.HexColor("#FCE4D6")))
            ts.append(("BACKGROUND", (5,i), (5,i), colors.HexColor("#FFF2CC")))
        else:
            ts.append(("BACKGROUND", (3,i), (3,i), colors.HexColor("#FFF2CC")))
    table.setStyle(TableStyle(ts))
    story.append(table)

    story.append(_pdf_p("2. 영역별 성취도", heading_style))
    area_headers = ["영역", "정답", "채점 문항", "정답률", "판정"]
    area_data = [[_pdf_p(x, header_style) for x in area_headers]]
    for a in result["areas"]:
        area_data.append([
            _pdf_p(a["area"], body), _pdf_p(a["correct"], body_center),
            _pdf_p(a["count"], body_center), _pdf_p(f'{a["rate"]}%', body_center),
            _pdf_p(a["judgment"], body_center),
        ])
    at = Table(area_data, colWidths=[usable*0.42, usable*0.12, usable*0.16, usable*0.15, usable*0.15], repeatRows=1)
    ats = [
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#1F4E79")),
        ("GRID", (0,0), (-1,-1), 0.25, colors.HexColor("#B7C9D6")),
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ("LEFTPADDING", (0,0), (-1,-1), 3),
        ("RIGHTPADDING", (0,0), (-1,-1), 3),
        ("TOPPADDING", (0,0), (-1,-1), 1.0),
        ("BOTTOMPADDING", (0,0), (-1,-1), 1.0),
    ]
    for i in range(1, len(area_data)):
        if i % 2 == 0:
            ats.append(("BACKGROUND", (0,i), (-1,i), colors.HexColor("#EEF3F8")))
    at.setStyle(TableStyle(ats))
    story.append(at)
    story.append(Spacer(1, 1.5))
    story.append(_pdf_p("※ 채점 문항이 1~2개인 영역은 표본이 적어 판정을 '참고'로 표시합니다.", small))

    story.append(_pdf_p("3. 오답 기반 진단", heading_style))
    diag_counts = defaultdict(int)
    for x in result["rows"]:
        if x["mark"] == "×" and x["diagnosis"]:
            diag_counts[x["diagnosis"]] += 1
    if diag_counts:
        for d, count in sorted(diag_counts.items(), key=lambda x: (-x[1], x[0]))[:8]:
            story.append(_pdf_p(f"• {d} - 관련 오답 {count}문항", body))
    else:
        story.append(_pdf_p("채점된 문항에서 별도 오답 진단 항목이 없습니다.", body))

    story.append(_pdf_p("4. 종합 학습 코멘트", heading_style))
    story.append(_pdf_p(comment_text(result), body))

    doc.build(story)
    return buf.getvalue()



def safe_name(s):
    s = re.sub(r'[\\/:*?"<>|]+', "_", s.strip())
    return s or "학생"


def reports_zip(reports: list[tuple[str, bytes]], exam: ExamAnalysis) -> bytes:
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z:
        for filename, data in reports:
            z.writestr(filename, data)
        z.writestr(
            "_시험분석_검수용.json",
            exam.model_dump_json(indent=2),
        )
    return out.getvalue()


st.set_page_config(
    page_title="성취도평가 리포트",
    page_icon="📘",
    layout="wide",
)

st.markdown("""
<style>
.block-container {max-width: 1100px; padding-top: 2rem; padding-bottom: 3rem;}
div[data-testid="stMetric"] {border: 1px solid #d8dee9; border-radius: 12px; padding: 12px;}
</style>
""", unsafe_allow_html=True)

st.title("📘 성취도평가 학생 리포트")
st.caption("시험지 PDF와 학생 답안만 올리면 문항 분석 → 채점 → 영역별 성취도 → Word 리포트까지 생성합니다.")

with st.sidebar:
    st.header("설정")
    default_key = ""
    try:
        default_key = st.secrets.get("OPENAI_API_KEY", "")
    except Exception:
        pass
    api_key = st.text_input(
        "OpenAI API 키",
        value=default_key,
        type="password",
        help="웹앱에 서버 비밀키가 설정되어 있으면 입력하지 않아도 됩니다.",
    )
    st.info("공식 정답표가 있으면 함께 올리는 것이 가장 정확합니다.")
    st.caption("학생 답안은 학생 1명당 파일 1개를 권장합니다. 여러 파일 동시 선택이 안 되면 ZIP 하나로 묶어 올리세요.")

col1, col2 = st.columns(2)
with col1:
    exam_file = st.file_uploader(
        "① 시험지 PDF",
        type=["pdf"],
        accept_multiple_files=False,
    )
with col2:
    answer_key = st.file_uploader(
        "② 공식 정답표/채점기준 (선택)",
        type=["pdf", "png", "jpg", "jpeg"],
        accept_multiple_files=False,
    )

student_files = st.file_uploader(
    "③ 학생 답안 여러 개",
    type=["pdf", "png", "jpg", "jpeg", "webp"],
    accept_multiple_files=True,
    help="여러 파일 선택이 어려운 기기에서는 아래 ZIP 업로드를 사용하세요.",
)

student_zip = st.file_uploader(
    "③-보조 학생 답안 ZIP 한꺼번에 올리기",
    type=["zip"],
    accept_multiple_files=False,
    help="태블릿/휴대폰에서 여러 파일 동시 선택이 안 되면, 학생 답안들을 ZIP 하나로 묶어 올리세요.",
)

# ZIP이 올라오면 안의 PDF/JPG/PNG/WEBP를 학생 파일 목록에 추가
if student_zip is not None:
    try:
        extracted = []
        with zipfile.ZipFile(io.BytesIO(student_zip.getvalue())) as zf:
            for info in zf.infolist():
                if info.is_dir():
                    continue
                base = os.path.basename(info.filename)
                if not base or base.startswith("."):
                    continue
                ext = os.path.splitext(base)[1].lower()
                if ext in [".pdf", ".png", ".jpg", ".jpeg", ".webp"]:
                    extracted.append(MemoryUploadedFile(base, zf.read(info)))
        if extracted:
            student_files = list(student_files or []) + extracted
            st.caption(f"ZIP에서 학생 답안 {len(extracted)}개를 불러왔습니다.")
        else:
            st.warning("ZIP 안에서 지원되는 학생 답안 파일(PDF/JPG/PNG/WEBP)을 찾지 못했습니다.")
    except zipfile.BadZipFile:
        st.error("올린 파일이 정상적인 ZIP 파일이 아닙니다.")

st.divider()

if "exam_analysis" not in st.session_state:
    st.session_state.exam_analysis = None
if "reports" not in st.session_state:
    st.session_state.reports = None
if "zip_bytes" not in st.session_state:
    st.session_state.zip_bytes = None

run = st.button("🚀 분석하고 Word 리포트 만들기", type="primary", use_container_width=True)

if run:
    if not api_key:
        st.error("OpenAI API 키가 필요합니다. 왼쪽 설정에서 입력하거나 서버 비밀키로 설정해 주세요.")
        st.stop()
    if exam_file is None:
        st.error("시험지 PDF를 올려 주세요.")
        st.stop()
    if not student_files:
        st.error("학생 답안 파일을 1개 이상 올려 주세요.")
        st.stop()

    progress = st.progress(0)
    status = st.empty()

    try:
        status.write("시험지를 문항별로 분석하고 있습니다...")
        exam = analyze_exam(api_key, exam_file, answer_key)
        st.session_state.exam_analysis = exam
        progress.progress(20)

        reports = []
        results = []
        count = len(student_files)

        for idx, f in enumerate(student_files, 1):
            status.write(f"학생 답안 채점 중: {f.name} ({idx}/{count})")
            grading = grade_student(api_key, exam, f)
            result = calculate(exam, grading)
            results.append(result)

            cls = safe_name(result["class"]) if result["class"] else ""
            name = safe_name(result["name"])
            base = f"{cls + '_' if cls else ''}{name}_리포트"
            reports.append((base + ".docx", make_docx(exam, result)))
            reports.append((base + ".pdf", make_pdf(exam, result)))

            progress.progress(20 + int(75 * idx / count))

        zip_bytes = reports_zip(reports, exam)
        st.session_state.reports = (reports, results)
        st.session_state.zip_bytes = zip_bytes
        progress.progress(100)
        status.success(f"완료: {len(results)}명의 Word + PDF 리포트를 만들었습니다.")

    except Exception as e:
        st.exception(e)

if st.session_state.exam_analysis is not None:
    exam = st.session_state.exam_analysis
    st.subheader("시험 분석 확인")
    c1, c2, c3 = st.columns(3)
    c1.metric("문항 수", exam.total_questions)
    c2.metric("시험지 표기 총점", exam.stated_total_points)
    ambiguous = sum(1 for q in exam.questions if q.answer_status != "certain")
    c3.metric("확인/제외 문항", ambiguous)

    with st.expander("문항 분석표 보기"):
        st.dataframe(
            [
                {
                    "번호": q.number,
                    "유형": q.question_type,
                    "배점": q.points,
                    "영역": q.area,
                    "출제 포인트": q.point,
                    "정답 상태": q.answer_status,
                    "분석 근거": q.evidence,
                }
                for q in exam.questions
            ],
            use_container_width=True,
            hide_index=True,
        )
        if exam.notes:
            st.caption(exam.notes)

if st.session_state.reports:
    reports, results = st.session_state.reports
    st.subheader("학생별 결과")
    st.dataframe(
        [
            {
                "학생": r["name"],
                "반": r["class"],
                "채점 문항": r["graded"],
                "정답": r["correct"],
                "오답": r["wrong"],
                "정답률": f'{r["rate"]}%',
                "채점 점수": f'{r["score"]}/{r["possible"]}',
                "확인 필요": r["confirm"],
                "채점 제외": r["excluded"],
            }
            for r in results
        ],
        use_container_width=True,
        hide_index=True,
    )

    st.download_button(
        "📦 모든 학생 Word + PDF 리포트 ZIP 다운로드",
        data=st.session_state.zip_bytes,
        file_name="학생_성취도평가_리포트.zip",
        mime="application/zip",
        type="primary",
        use_container_width=True,
    )

    with st.expander("학생별 Word / PDF 파일 따로 받기"):
        for filename, data in reports:
            mime = (
                "application/pdf"
                if filename.lower().endswith(".pdf")
                else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.download_button(
                filename,
                data=data,
                file_name=filename,
                mime=mime,
                key="download_" + filename,
            )

st.divider()
st.caption(
    "정확성 원칙: 공식 정답표를 우선하며, 모호하거나 판독 불가한 문항은 임의로 채점하지 않고 확인 필요/채점 제외로 처리합니다."
)
